from __future__ import annotations

import base64
import io
import json
import os
import re
import sys
import threading
import time
import uuid
import webbrowser
from collections import OrderedDict
from pathlib import Path
from typing import Any, Iterable

import fitz
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from flask import Flask, jsonify, render_template, request, send_file
from PIL import Image, ImageColor, ImageDraw, ImageFont
from werkzeug.exceptions import RequestEntityTooLarge


APP_VERSION = "1.0.0"
APP_ROOT = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
MAX_UPLOAD_BYTES = 50 * 1024 * 1024
MAX_CACHED_DOCUMENTS = 8
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm", ".png", ".jpg", ".jpeg", ".webp"}
PAGE_SIZE = (1240, 1754)  # A4 at approximately 150 dpi

app = Flask(
    __name__,
    static_folder=str(APP_ROOT / "static"),
    template_folder=str(APP_ROOT / "templates"),
)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_BYTES

# This app is intentionally local-first. Imported page previews stay in memory and
# are dropped as older documents are opened.
_documents: OrderedDict[str, dict[str, Any]] = OrderedDict()


def _font_file_path(
    bold: bool = False,
    italic: bool = False,
    family: str = "sans",
) -> Path | None:
    family = family.lower()
    if "calibri" in family:
        windows_name = "calibriz.ttf" if bold and italic else "calibrib.ttf" if bold else "calibrii.ttf" if italic else "calibri.ttf"
        linux_name = "DejaVuSans-BoldOblique.ttf" if bold and italic else "DejaVuSans-Bold.ttf" if bold else "DejaVuSans-Oblique.ttf" if italic else "DejaVuSans.ttf"
        linux_folder = "dejavu"
    elif "mono" in family or "consol" in family or "courier" in family:
        windows_name = "consolaz.ttf" if bold and italic else "consolab.ttf" if bold else "consolai.ttf" if italic else "consola.ttf"
        linux_name = "DejaVuSansMono-BoldOblique.ttf" if bold and italic else "DejaVuSansMono-Bold.ttf" if bold else "DejaVuSansMono-Oblique.ttf" if italic else "DejaVuSansMono.ttf"
        linux_folder = "dejavu"
    elif ("serif" in family and "sans-serif" not in family) or "times" in family or "georgia" in family:
        windows_name = "timesbi.ttf" if bold and italic else "timesbd.ttf" if bold else "timesi.ttf" if italic else "times.ttf"
        linux_name = "DejaVuSerif-BoldItalic.ttf" if bold and italic else "DejaVuSerif-Bold.ttf" if bold else "DejaVuSerif-Italic.ttf" if italic else "DejaVuSerif.ttf"
        linux_folder = "dejavu"
    else:
        windows_name = "arialbi.ttf" if bold and italic else "arialbd.ttf" if bold else "ariali.ttf" if italic else "arial.ttf"
        linux_name = "DejaVuSans-BoldOblique.ttf" if bold and italic else "DejaVuSans-Bold.ttf" if bold else "DejaVuSans-Oblique.ttf" if italic else "DejaVuSans.ttf"
        linux_folder = "dejavu"
    candidates = [
        Path("C:/Windows/Fonts") / windows_name,
        Path("/usr/share/fonts/truetype") / linux_folder / linux_name,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _font(
    size: int,
    bold: bool = False,
    italic: bool = False,
    family: str = "sans",
) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_path = _font_file_path(bold, italic, family)
    if font_path:
        return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def _png_bytes(image: Image.Image) -> bytes:
    output = io.BytesIO()
    image.convert("RGB").save(output, format="PNG", optimize=True)
    return output.getvalue()


def _remember_document(name: str, pages: list[bytes], dimensions: list[tuple[int, int]]) -> str:
    document_id = uuid.uuid4().hex
    _documents[document_id] = {
        "name": name,
        "pages": pages,
        "dimensions": dimensions,
        "created": time.time(),
    }
    _documents.move_to_end(document_id)
    while len(_documents) > MAX_CACHED_DOCUMENTS:
        _documents.popitem(last=False)
    return document_id


def _wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    if not text:
        return [""]
    words = text.replace("\t", "    ").split(" ")
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if not line or draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines or [""]


def _iter_docx_blocks(document: Document) -> Iterable[dict[str, Any]]:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            style = (paragraph.style.name if paragraph.style else "") or ""
            yield {"type": "paragraph", "text": paragraph.text, "style": style}
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            yield {
                "type": "table",
                "rows": [[cell.text.strip() for cell in row.cells] for row in table.rows],
            }


def _text_annotation(
    text: str,
    x: float,
    y: float,
    width: float,
    height: float,
    font_size: float,
    color: str = "#172033",
    bold: bool = False,
    italic: bool = False,
    font_family: str = "Arial, sans-serif",
) -> dict[str, Any]:
    return {
        "id": uuid.uuid4().hex,
        "type": "text",
        "source": "imported",
        "x": round(x, 2),
        "y": round(y, 2),
        "width": round(max(width, font_size), 2),
        "height": round(max(height, font_size * 1.25), 2),
        "text": text,
        "fontSize": round(font_size, 2),
        "fontFamily": font_family,
        "color": color,
        "opacity": 1,
        "bold": bold,
        "italic": italic,
        "cover": False,
    }


def _render_blocks(
    blocks: Iterable[dict[str, Any]],
) -> tuple[list[bytes], list[tuple[int, int]], list[list[dict[str, Any]]]]:
    page_width, page_height = PAGE_SIZE
    margin_x, margin_y = 105, 105
    usable_width = page_width - margin_x * 2
    images: list[Image.Image] = []
    page_annotations: list[list[dict[str, Any]]] = []
    image: Image.Image
    draw: ImageDraw.ImageDraw
    y: int

    def new_page() -> None:
        nonlocal image, draw, y
        image = Image.new("RGB", PAGE_SIZE, "white")
        draw = ImageDraw.Draw(image)
        y = margin_y
        images.append(image)
        page_annotations.append([])

    def ensure_space(required: int) -> None:
        if y + required > page_height - margin_y:
            new_page()

    new_page()
    for block in blocks:
        if block.get("type") == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            cell_width = usable_width // max(column_count, 1)
            table_font = _font(21)
            for row in rows:
                wrapped_cells = [
                    _wrap_text(draw, str(row[index]) if index < len(row) else "", table_font, cell_width - 20)
                    for index in range(column_count)
                ]
                row_height = max(48, max(len(lines) for lines in wrapped_cells) * 29 + 18)
                ensure_space(row_height)
                for index, lines in enumerate(wrapped_cells):
                    left = margin_x + index * cell_width
                    right = margin_x + (index + 1) * cell_width
                    draw.rectangle((left, y, right, y + row_height), outline="#cbd5e1", width=2)
                    cell_text = str(row[index]) if index < len(row) else ""
                    if cell_text:
                        page_annotations[-1].append(_text_annotation(
                            cell_text,
                            left + 10,
                            y + 9,
                            cell_width - 20,
                            row_height - 18,
                            21,
                            color="#1e293b",
                        ))
                y += row_height
            y += 24
            continue

        text = str(block.get("text", ""))
        style = str(block.get("style", "")).lower()
        if not text.strip():
            y += 18
            continue
        if "title" in style:
            size, bold, spacing = 45, True, 28
        elif "heading 1" in style or block.get("level") == 1:
            size, bold, spacing = 36, True, 22
        elif "heading 2" in style or block.get("level") == 2:
            size, bold, spacing = 30, True, 18
        elif block.get("bullet"):
            size, bold, spacing = 23, False, 12
            text = f"•  {text}"
        else:
            size, bold, spacing = 23, False, 14
        font = _font(size, bold)
        line_height = int(size * 1.45)
        lines = _wrap_text(draw, text, font, usable_width)
        required = len(lines) * line_height + spacing
        ensure_space(required)
        page_annotations[-1].append(_text_annotation(
            text,
            margin_x,
            y,
            usable_width,
            len(lines) * line_height,
            size,
            bold=bold,
        ))
        y += len(lines) * line_height
        y += spacing

    return [_png_bytes(page) for page in images], [PAGE_SIZE for _ in images], page_annotations


def _pdf_text_family(flags: int, embedded_name: str = "") -> str:
    normalized = embedded_name.split("+")[-1].lower().replace("-", " ")
    if "calibri" in normalized:
        return "Calibri, Arial, sans-serif"
    if "aptos" in normalized:
        return "Aptos, Calibri, Arial, sans-serif"
    if "cambria" in normalized:
        return "Cambria, Georgia, serif"
    if "georgia" in normalized:
        return "Georgia, Times New Roman, serif"
    if "times" in normalized:
        return "Times New Roman, Times, serif"
    if "consol" in normalized or "courier" in normalized:
        return "Consolas, Courier New, monospace"
    if "arial" in normalized or "helvetica" in normalized:
        return "Arial, Helvetica, sans-serif"
    if flags & 8:
        return "Consolas, Courier New, monospace"
    if flags & 4:
        return "Times New Roman, Times, serif"
    return "Arial, Helvetica, sans-serif"


def _pdf_text_color(value: int) -> str:
    return f"#{int(value) & 0xFFFFFF:06x}"


def _editable_pdf_pages(
    payload: bytes,
) -> tuple[list[bytes], list[tuple[int, int]], list[list[dict[str, Any]]]]:
    source = fitz.open(stream=payload, filetype="pdf")
    if source.page_count == 0:
        source.close()
        raise ValueError("This PDF has no pages.")

    pages: list[bytes] = []
    dimensions: list[tuple[int, int]] = []
    page_annotations: list[list[dict[str, Any]]] = []
    render_scale = 1.6
    matrix = fitz.Matrix(render_scale, render_scale)

    for pdf_page in source:
        annotations: list[dict[str, Any]] = []
        redact_rects: list[fitz.Rect] = []
        text_data = pdf_page.get_text("dict")
        rotation_matrix = pdf_page.rotation_matrix
        for block in text_data.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = str(span.get("text", ""))
                    if not text:
                        continue
                    source_rect = fitz.Rect(span["bbox"])
                    display_rect = source_rect * rotation_matrix
                    display_rect.normalize()
                    display_origin = fitz.Point(span.get("origin", (source_rect.x0, source_rect.y1))) * rotation_matrix
                    flags = int(span.get("flags", 0))
                    alpha = max(0, min(int(span.get("alpha", 255)), 255)) / 255
                    annotation = _text_annotation(
                        text=text,
                        x=display_rect.x0 * render_scale,
                        y=display_rect.y0 * render_scale,
                        width=display_rect.width * render_scale + 2,
                        height=display_rect.height * render_scale,
                        font_size=float(span.get("size", 11)) * render_scale,
                        color=_pdf_text_color(int(span.get("color", 0))),
                        bold=bool(flags & 16),
                        italic=bool(flags & 2),
                        font_family=_pdf_text_family(flags, str(span.get("font", ""))),
                    )
                    annotation["opacity"] = round(alpha, 3)
                    annotation["baselineY"] = round(display_origin.y * render_scale, 2)
                    annotation["noWrap"] = True
                    annotations.append(annotation)
                    redact_rects.append(source_rect)

        for rect in redact_rects:
            pdf_page.add_redact_annot(rect, fill=None, cross_out=False)
        if redact_rects:
            pdf_page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )

        pixmap = pdf_page.get_pixmap(matrix=matrix, alpha=False)
        pages.append(pixmap.tobytes("png"))
        dimensions.append((pixmap.width, pixmap.height))
        page_annotations.append(annotations)

    source.close()
    return pages, dimensions, page_annotations


def _document_to_pages(
    filename: str,
    payload: bytes,
) -> tuple[list[bytes], list[tuple[int, int]], list[list[dict[str, Any]]]]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _editable_pdf_pages(payload)

    if extension == ".docx":
        document = Document(io.BytesIO(payload))
        return _render_blocks(_iter_docx_blocks(document))

    if extension in {".txt", ".md", ".markdown", ".html", ".htm"}:
        raw = payload.decode("utf-8", errors="replace")
        blocks: list[dict[str, Any]] = []
        if extension in {".html", ".htm"}:
            soup = BeautifulSoup(raw, "html.parser")
            for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
                text = element.get_text(" ", strip=True)
                if text:
                    blocks.append({
                        "type": "paragraph",
                        "text": text,
                        "level": 1 if element.name == "h1" else 2 if element.name in {"h2", "h3"} else None,
                        "bullet": element.name == "li",
                    })
        else:
            for line in raw.splitlines():
                stripped = line.strip()
                level = 1 if stripped.startswith("# ") else 2 if stripped.startswith(("## ", "### ")) else None
                bullet = bool(re.match(r"^[-*+]\s+", stripped))
                stripped = re.sub(r"^#{1,3}\s+", "", stripped)
                stripped = re.sub(r"^[-*+]\s+", "", stripped)
                blocks.append({"type": "paragraph", "text": stripped, "level": level, "bullet": bullet})
        return _render_blocks(blocks)

    if extension in {".png", ".jpg", ".jpeg", ".webp"}:
        image = Image.open(io.BytesIO(payload)).convert("RGB")
        image.thumbnail((2400, 2400), Image.Resampling.LANCZOS)
        return [_png_bytes(image)], [image.size], [[]]

    raise ValueError("That file format is not supported.")


def _parse_color(value: str, alpha: int = 255) -> tuple[int, int, int, int]:
    try:
        red, green, blue = ImageColor.getrgb(value or "#111827")[:3]
    except ValueError:
        red, green, blue = (17, 24, 39)
    return red, green, blue, max(0, min(alpha, 255))


def _resolve_background(page: dict[str, Any]) -> Image.Image:
    data_url = page.get("imageData")
    if isinstance(data_url, str) and data_url.startswith("data:image/"):
        encoded = data_url.split(",", 1)[1]
        return Image.open(io.BytesIO(base64.b64decode(encoded))).convert("RGBA")

    document_id = str(page.get("documentId", ""))
    page_index = int(page.get("serverIndex", -1))
    cached = _documents.get(document_id)
    if not cached or page_index < 0 or page_index >= len(cached["pages"]):
        raise ValueError("A page background is no longer available. Reopen the source file and try again.")
    return Image.open(io.BytesIO(cached["pages"][page_index])).convert("RGBA")


def _compose_page(page: dict[str, Any]) -> Image.Image:
    image = _resolve_background(page)
    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    scale_x = image.width / max(float(page.get("width", image.width)), 1)
    scale_y = image.height / max(float(page.get("height", image.height)), 1)

    for annotation in page.get("annotations", []):
        annotation_type = annotation.get("type")
        opacity = float(annotation.get("opacity", 1))
        alpha = int(max(0, min(opacity, 1)) * 255)
        color = _parse_color(str(annotation.get("color", "#111827")), alpha)
        if annotation_type == "text":
            text = str(annotation.get("text", ""))
            if not text:
                continue
            x = int(float(annotation.get("x", 0)) * scale_x)
            y = int(float(annotation.get("y", 0)) * scale_y)
            width = int(float(annotation.get("width", 280)) * scale_x)
            font_size = max(8, int(float(annotation.get("fontSize", 24)) * min(scale_x, scale_y)))
            font = _font(
                font_size,
                bool(annotation.get("bold")),
                bool(annotation.get("italic")),
                str(annotation.get("fontFamily", "sans")),
            )
            line_height = int(font_size * 1.25)
            if annotation.get("cover"):
                height = int(float(annotation.get("height", line_height * 1.25)) * scale_y)
                draw.rectangle((x, y, x + width, y + height), fill=(255, 255, 255, 255))
            if not annotation.get("noWrap") and annotation.get("source") != "imported":
                x += round(5 * scale_x)
                y += round(3 * scale_y)
                width = max(width - round(10 * scale_x), 1)
            for line in text.splitlines() or [text]:
                rendered_lines = [line] if annotation.get("noWrap") else _wrap_text(draw, line, font, width)
                for wrapped in rendered_lines:
                    if annotation.get("noWrap") and wrapped:
                        draw.text((x, y), wrapped, fill=color, font=font, anchor="lt")
                    else:
                        draw.text((x, y), wrapped, fill=color, font=font, anchor="lt")
                    y += line_height
        elif annotation_type == "stroke":
            points = annotation.get("points") or []
            if len(points) < 2:
                continue
            scaled_points = [
                (float(point[0]) * scale_x, float(point[1]) * scale_y)
                for point in points
            ]
            line_width = max(1, int(float(annotation.get("lineWidth", 4)) * min(scale_x, scale_y)))
            draw.line(scaled_points, fill=color, width=line_width, joint="curve")

    return Image.alpha_composite(image, overlay).convert("RGB")


def _pdf_base_font_name(bold: bool, italic: bool) -> str:
    if bold and italic:
        return "hebi"
    if bold:
        return "hebo"
    if italic:
        return "heit"
    return "helv"


def _register_pdf_font(
    pdf_page: fitz.Page,
    annotation: dict[str, Any],
    font_cache: dict[tuple[str, bool, bool], str],
) -> str:
    family = str(annotation.get("fontFamily", "sans"))
    bold = bool(annotation.get("bold"))
    italic = bool(annotation.get("italic"))
    key = (family, bold, italic)
    if key in font_cache:
        return font_cache[key]

    font_path = _font_file_path(bold, italic, family)
    if font_path:
        font_name = f"PF{len(font_cache) + 1}"
        pdf_page.insert_font(fontname=font_name, fontfile=str(font_path))
    else:
        font_name = _pdf_base_font_name(bold, italic)
    font_cache[key] = font_name
    return font_name


def _build_selectable_pdf(pages: list[dict[str, Any]]) -> bytes:
    pdf = fitz.open()
    points_per_pixel = 72 / 150

    for page_spec in pages:
        background = _resolve_background(page_spec).convert("RGB")
        background_buffer = io.BytesIO()
        background.save(background_buffer, format="JPEG", quality=94, optimize=True)
        source_width = max(float(page_spec.get("width", background.width)), 1)
        source_height = max(float(page_spec.get("height", background.height)), 1)
        page_width = source_width * points_per_pixel
        page_height = source_height * points_per_pixel
        scale_x = page_width / source_width
        scale_y = page_height / source_height
        pdf_page = pdf.new_page(width=page_width, height=page_height)
        pdf_page.insert_image(pdf_page.rect, stream=background_buffer.getvalue())
        font_cache: dict[tuple[str, bool, bool], str] = {}

        for annotation in page_spec.get("annotations", []):
            annotation_type = annotation.get("type")
            opacity = max(0, min(float(annotation.get("opacity", 1)), 1))
            red, green, blue, _alpha = _parse_color(str(annotation.get("color", "#111827")))
            color = (red / 255, green / 255, blue / 255)

            if annotation_type == "text":
                text = str(annotation.get("text", ""))
                if not text:
                    continue
                x = float(annotation.get("x", 0)) * scale_x
                y = float(annotation.get("y", 0)) * scale_y
                width = max(float(annotation.get("width", 280)) * scale_x, 1)
                height = max(float(annotation.get("height", 50)) * scale_y, 1)
                font_size = max(float(annotation.get("fontSize", 24)) * min(scale_x, scale_y), 1)
                if annotation.get("cover"):
                    pdf_page.draw_rect(
                        fitz.Rect(x, y, x + width, y + height),
                        color=(1, 1, 1),
                        fill=(1, 1, 1),
                        overlay=True,
                    )
                if not annotation.get("noWrap") and annotation.get("source") != "imported":
                    x += 5 * scale_x
                    y += 3 * scale_y
                    width = max(width - 10 * scale_x, 1)
                if not annotation.get("noWrap"):
                    # PyMuPDF text boxes reserve ascender space above visible glyphs;
                    # offset it so their visual top matches the browser / Pillow editor.
                    y -= font_size * 0.21
                font_name = _register_pdf_font(pdf_page, annotation, font_cache)
                if annotation.get("noWrap"):
                    baseline = float(annotation.get("baselineY", annotation.get("y", 0) + annotation.get("fontSize", 24) * 0.9)) * scale_y
                    pdf_page.insert_text(
                        fitz.Point(x, baseline),
                        text,
                        fontname=font_name,
                        fontsize=font_size,
                        color=color,
                        fill_opacity=opacity,
                        lineheight=1.15,
                        overlay=True,
                    )
                else:
                    line_count = max(len(text.splitlines()), 1)
                    textbox_height = max(
                        height + font_size,
                        line_count * font_size * 1.4 + font_size * 0.5,
                    )
                    pdf_page.insert_textbox(
                        fitz.Rect(x, y, x + width, min(y + textbox_height, page_height)),
                        text,
                        fontname=font_name,
                        fontsize=font_size,
                        lineheight=1.25,
                        color=color,
                        fill_opacity=opacity,
                        overlay=True,
                    )
            elif annotation_type == "stroke":
                points = annotation.get("points") or []
                if len(points) < 2:
                    continue
                pdf_points = [
                    fitz.Point(float(point[0]) * scale_x, float(point[1]) * scale_y)
                    for point in points
                ]
                pdf_page.draw_polyline(
                    pdf_points,
                    color=color,
                    width=max(float(annotation.get("lineWidth", 4)) * min(scale_x, scale_y), 0.5),
                    lineCap=1,
                    lineJoin=1,
                    stroke_opacity=opacity,
                    overlay=True,
                )

    result = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return result


def _safe_download_name(value: str, extension: str) -> str:
    stem = Path(value or "edited-document").stem
    stem = re.sub(r"[^A-Za-z0-9 _.-]", "", stem).strip(" .") or "edited-document"
    return f"{stem}.{extension}"


@app.get("/")
def index() -> str:
    return render_template("index.html")


@app.post("/api/import")
def import_document():
    upload = request.files.get("file")
    if upload is None or not upload.filename:
        return jsonify({"error": "Choose a file to open."}), 400
    extension = Path(upload.filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        return jsonify({"error": "Unsupported file. Open a PDF, DOCX, text, Markdown, HTML, PNG, JPG, or WebP file."}), 415
    payload = upload.read()
    if not payload:
        return jsonify({"error": "The selected file is empty."}), 400
    try:
        pages, dimensions, annotations = _document_to_pages(upload.filename, payload)
    except Exception as exc:
        app.logger.exception("Import failed")
        return jsonify({"error": f"Could not open this file: {exc}"}), 422

    document_id = _remember_document(upload.filename, pages, dimensions)
    return jsonify({
        "documentId": document_id,
        "name": Path(upload.filename).stem,
        "sourceFormat": extension.lstrip("."),
        "pages": [
            {
                "width": width,
                "height": height,
                "serverIndex": index,
                "imageUrl": f"/api/documents/{document_id}/pages/{index}",
                "annotations": annotations[index],
            }
            for index, (width, height) in enumerate(dimensions)
        ],
        "editableTextCount": sum(len(items) for items in annotations),
    })


@app.get("/api/documents/<document_id>/pages/<int:page_index>")
def document_page(document_id: str, page_index: int):
    cached = _documents.get(document_id)
    if not cached or page_index < 0 or page_index >= len(cached["pages"]):
        return jsonify({"error": "Page not found."}), 404
    _documents.move_to_end(document_id)
    return send_file(io.BytesIO(cached["pages"][page_index]), mimetype="image/png", max_age=3600)


@app.post("/api/export")
def export_document():
    payload = request.get_json(silent=True) or {}
    pages = payload.get("pages") or []
    export_format = str(payload.get("format", "pdf")).lower()
    if not pages:
        return jsonify({"error": "There are no pages to export."}), 400
    if export_format not in {"pdf", "docx", "png", "jpg"}:
        return jsonify({"error": "Unsupported export format."}), 400
    try:
        if export_format in {"png", "jpg"}:
            page_index = max(0, min(int(payload.get("pageIndex", 0)), len(pages) - 1))
            image = _compose_page(pages[page_index])
            output = io.BytesIO()
            if export_format == "jpg":
                image.save(output, format="JPEG", quality=94, optimize=True)
                mimetype = "image/jpeg"
            else:
                image.save(output, format="PNG", optimize=True)
                mimetype = "image/png"
            output.seek(0)
            return send_file(output, mimetype=mimetype, as_attachment=True, download_name=_safe_download_name(payload.get("name", ""), export_format))

        output = io.BytesIO()
        if export_format == "pdf":
            output.write(_build_selectable_pdf(pages))
            mimetype = "application/pdf"
        else:
            composed = [_compose_page(page) for page in pages]
            document = Document()
            section = document.sections[0]
            section.top_margin = Inches(0.35)
            section.bottom_margin = Inches(0.35)
            section.left_margin = Inches(0.35)
            section.right_margin = Inches(0.35)
            for index, image in enumerate(composed):
                if index:
                    document.add_section(WD_SECTION.NEW_PAGE)
                image_buffer = io.BytesIO()
                image.save(image_buffer, format="PNG", optimize=True)
                max_width, max_height = 7.55, 10.9
                aspect = image.width / image.height
                width = min(max_width, max_height * aspect)
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.add_run().add_picture(image_buffer, width=Inches(width))
            document.save(output)
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        output.seek(0)
        return send_file(output, mimetype=mimetype, as_attachment=True, download_name=_safe_download_name(payload.get("name", ""), export_format))
    except Exception as exc:
        app.logger.exception("Export failed")
        return jsonify({"error": f"Could not export this document: {exc}"}), 422


@app.errorhandler(RequestEntityTooLarge)
def upload_too_large(_error: RequestEntityTooLarge):
    return jsonify({"error": "That file is larger than the 50 MB limit."}), 413


@app.get("/api/health")
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    try:
        port = int(os.environ.get("PDFEDITORATHOME_PORT", "5050"))
    except ValueError:
        port = 5050
    if not 1 <= port <= 65535:
        port = 5050
    address = f"http://127.0.0.1:{port}"
    if getattr(sys, "frozen", False) and os.environ.get("PDFEDITORATHOME_NO_BROWSER") != "1":
        browser_timer = threading.Timer(1.0, webbrowser.open, args=(address,))
        browser_timer.daemon = True
        browser_timer.start()
    print(f"PDFeditorAthome v{APP_VERSION} is running at {address}")
    print("Keep this window open while using PDFeditorAthome. Close it to stop the app.")
    app.run(host="127.0.0.1", port=port, debug=False, use_reloader=False)
