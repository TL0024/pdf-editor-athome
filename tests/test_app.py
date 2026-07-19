import base64
import io
from typing import Any, cast

import fitz
from docx import Document
from flask.testing import FlaskClient
from PIL import Image, ImageChops

from app import app


def client() -> FlaskClient:
    app.config.update(TESTING=True)
    return app.test_client()


def import_text_page(test_client: FlaskClient) -> dict[str, Any]:
    response = test_client.post(
        "/api/import",
        data={"file": (io.BytesIO(b"# Sample document\n\nA paragraph for export."), "sample.md")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    return response.get_json()


def export_payload(imported: dict[str, Any], export_format: str) -> dict[str, Any]:
    page = imported["pages"][0]
    return {
        "name": "Edited sample",
        "format": export_format,
        "pageIndex": 0,
        "pages": [
            {
                "width": page["width"],
                "height": page["height"],
                "documentId": imported["documentId"],
                "serverIndex": page["serverIndex"],
                "annotations": [
                    *page.get("annotations", []),
                    {
                        "id": "text-1",
                        "type": "text",
                        "x": 120,
                        "y": 300,
                        "width": 400,
                        "height": 60,
                        "text": "Editable note",
                        "fontSize": 28,
                        "color": "#4f46e5",
                        "opacity": 1,
                        "bold": True,
                    },
                    {
                        "id": "stroke-1",
                        "type": "stroke",
                        "points": [[100, 420], [220, 460], [370, 430]],
                        "lineWidth": 8,
                        "color": "#e11d48",
                        "opacity": 0.8,
                    },
                ],
            }
        ],
    }


def test_home_and_health() -> None:
    test_client = client()
    home = test_client.get("/")
    assert home.status_code == 200
    assert b"PDFeditorAthome" in home.data
    assert test_client.get("/api/health").get_json() == {"status": "ok"}


def test_import_text_and_serve_page() -> None:
    test_client = client()
    imported = import_text_page(test_client)
    assert imported["sourceFormat"] == "md"
    assert len(imported["pages"]) == 1
    assert imported["editableTextCount"] == 2
    assert {item["text"] for item in imported["pages"][0]["annotations"]} == {
        "Sample document",
        "A paragraph for export.",
    }
    page_response = test_client.get(imported["pages"][0]["imageUrl"])
    assert page_response.status_code == 200
    image = Image.open(io.BytesIO(page_response.data))
    assert image.size == (1240, 1754)


def test_import_pdf() -> None:
    source = fitz.open()
    page = source.new_page(width=300, height=500)
    page.draw_rect(page.rect, color=(0.8, 0.9, 1), fill=(0.8, 0.9, 1))
    page.insert_text((50, 80), "PDF import", fontsize=18, color=(0.1, 0.2, 0.6))
    pdf_bytes = source.tobytes()
    source.close()

    response = client().post(
        "/api/import",
        data={"file": (io.BytesIO(pdf_bytes), "source.pdf")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    data = response.get_json()
    assert data["sourceFormat"] == "pdf"
    assert len(data["pages"]) == 1
    assert data["pages"][0]["height"] > data["pages"][0]["width"]
    assert data["editableTextCount"] == 1
    annotation = data["pages"][0]["annotations"][0]
    assert annotation["text"] == "PDF import"
    assert annotation["type"] == "text"
    assert annotation["source"] == "imported"
    assert annotation["noWrap"] is True
    assert annotation["baselineY"] > annotation["y"]
    assert "Arial" in annotation["fontFamily"]

    background_response = client().get(data["pages"][0]["imageUrl"])
    background = Image.open(io.BytesIO(background_response.data)).convert("RGB")
    left, top = int(annotation["x"]), int(annotation["y"])
    right = min(background.width, int(annotation["x"] + annotation["width"]))
    bottom = min(background.height, int(annotation["y"] + annotation["height"]))
    pixels = cast(list[tuple[int, int, int]], list(background.crop((left, top, right, bottom)).get_flattened_data()))
    assert pixels
    assert min(min(pixel) for pixel in pixels) > 180


def test_pdf_styled_runs_preserve_outer_spaces() -> None:
    source = fitz.open()
    page = source.new_page(width=500, height=200)
    x, baseline = 40, 80
    runs = [
        ("Regular trailing  ", "helv"),
        ("  Bold separated  ", "hebo"),
        ("  end", "helv"),
    ]
    for text, font_name in runs:
        page.insert_text((x, baseline), text, fontname=font_name, fontsize=16)
        x += fitz.get_text_length(text, fontname=font_name, fontsize=16)
    pdf_bytes = source.tobytes()
    source.close()

    response = client().post(
        "/api/import",
        data={"file": (io.BytesIO(pdf_bytes), "spaces.pdf")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    annotations = response.get_json()["pages"][0]["annotations"]
    assert [annotation["text"] for annotation in annotations] == [text for text, _font in runs]


def test_import_docx_text_is_editable() -> None:
    source = io.BytesIO()
    document = Document()
    document.add_heading("Editable Word heading", level=1)
    document.add_paragraph("Editable Word paragraph")
    document.save(source)
    source.seek(0)

    response = client().post(
        "/api/import",
        data={"file": (source, "editable.docx")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    data = response.get_json()
    assert data["editableTextCount"] == 2
    assert [item["text"] for item in data["pages"][0]["annotations"]] == [
        "Editable Word heading",
        "Editable Word paragraph",
    ]


def test_export_pdf_png_and_docx() -> None:
    test_client = client()
    imported = import_text_page(test_client)

    pdf_response = test_client.post("/api/export", json=export_payload(imported, "pdf"))
    assert pdf_response.status_code == 200
    exported_pdf = fitz.open(stream=pdf_response.data, filetype="pdf")
    assert exported_pdf.page_count == 1
    selectable_text = exported_pdf[0].get_text().replace("\xa0", " ")
    assert "Sample document" in selectable_text
    assert "A paragraph for export." in selectable_text
    assert "Editable note" in selectable_text
    assert len(exported_pdf[0].get_drawings()) >= 1
    exported_pdf.close()

    png_response = test_client.post("/api/export", json=export_payload(imported, "png"))
    assert png_response.status_code == 200
    exported_png = Image.open(io.BytesIO(png_response.data))
    assert exported_png.size == (1240, 1754)

    docx_response = test_client.post("/api/export", json=export_payload(imported, "docx"))
    assert docx_response.status_code == 200
    exported_docx = Document(io.BytesIO(docx_response.data))
    assert len(exported_docx.inline_shapes) == 1


def test_added_text_alignment_matches_pdf_and_image_exports() -> None:
    background = Image.new("RGB", (600, 400), "white")
    background_buffer = io.BytesIO()
    background.save(background_buffer, format="PNG")
    image_data = "data:image/png;base64," + base64.b64encode(background_buffer.getvalue()).decode()
    annotation = {
        "id": "alignment-text",
        "type": "text",
        "x": 200,
        "y": 150,
        "width": 220,
        "height": 60,
        "text": "Hello alignment",
        "fontSize": 24,
        "fontFamily": "Arial, sans-serif",
        "color": "#000000",
        "opacity": 1,
        "bold": False,
        "italic": False,
        "cover": False,
    }
    payload = {
        "name": "alignment",
        "pageIndex": 0,
        "pages": [{"width": 600, "height": 400, "imageData": image_data, "annotations": [annotation]}],
    }
    test_client = client()
    png_response = test_client.post("/api/export", json={**payload, "format": "png"})
    pdf_response = test_client.post("/api/export", json={**payload, "format": "pdf"})
    assert png_response.status_code == 200
    assert pdf_response.status_code == 200

    png_image = Image.open(io.BytesIO(png_response.data)).convert("RGB")
    pdf = fitz.open(stream=pdf_response.data, filetype="pdf")
    pixmap = pdf[0].get_pixmap(matrix=fitz.Matrix(150 / 72, 150 / 72), alpha=False)
    pdf_image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
    pdf.close()
    white = Image.new("RGB", png_image.size, "white")
    png_box = ImageChops.difference(png_image, white).getbbox()
    pdf_box = ImageChops.difference(pdf_image, white).getbbox()
    assert png_box and pdf_box
    assert abs(pdf_box[0] - png_box[0]) <= 2
    assert abs(pdf_box[1] - png_box[1]) <= 2


def test_reject_unsupported_extension() -> None:
    response = client().post(
        "/api/import",
        data={"file": (io.BytesIO(b"not a presentation"), "slides.ppt")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 415
    assert "Unsupported" in response.get_json()["error"]
